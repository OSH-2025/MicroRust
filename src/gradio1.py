import re
import requests
import gradio as gr
import time
import os
from typing import Optional, Tuple
from datetime import datetime
from PyPDF2 import PdfReader
import docx
from io import BytesIO

# ==================== 配置 ====================
IPFS_API_URL = "http://192.168.67.128:5001"
DEEPSEEK_API_KEY = "sk-5c343e8522ef4787bcd862aa005af5b4"
MAX_FILE_SIZE = 10 * 1024 * 1024
CHUNK_SIZE = 1024 * 512
MAX_RETRIES = 3

# 存储已上传文件 {CID: (文件名, 上传时间, 文件类型)}
uploaded_files = {}

# 防抖相关变量
last_input_time = 0

# ==================== 提取文本用于预览 ====================
def extract_text_from_file(filepath: str) -> Tuple[str, str]:
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == ".pdf":
            reader = PdfReader(filepath)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text.strip(), "pdf"
        elif ext == ".docx":
            doc = docx.Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            return text.strip(), "docx"
        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read(), "txt"
        else:
            return "⚠️ 不支持的文件类型", "unknown"
    except Exception as e:
        return f"⚠️ 无法解析文件内容: {str(e)}", "error"

# ==================== 增强版CID匹配 ====================
def match_cid(msg: str) -> Tuple[Optional[str], str]:
    if not msg:
        return None, "🟡 未检测到文件引用"
    
    cid_pattern = r"\b(Qm[1-9A-HJ-NP-Za-km-z]{44}|bafy[0-9A-Za-z]{44,})\b"
    matches = re.findall(cid_pattern, msg)
    
    if matches:
        for cid in matches:
            if cid in uploaded_files:
                filename = uploaded_files[cid][0]
                return cid, f"✅ 已关联文件: {filename[:15]}..."
        return None, f"🟠 检测到CID但未上传: {matches[0][:6]}..."
    
    return None, "🟡 未检测到文件引用"

# ==================== 防抖函数 ====================
def debounced_match_cid(message):
    global last_input_time
    current_time = time.time() * 1000
    if current_time - last_input_time < 500:
        return "🟡 正在输入..."
    last_input_time = current_time
    return match_cid(message)[1]

# ==================== 文件上传处理 ====================
def handle_upload(file) -> str:
    if not file:
        return "⚠️ 请先选择文件"
    
    try:
        # 上传原始文件到 IPFS
        with open(file.name, "rb") as f:
            response = requests.post(
                f"{IPFS_API_URL}/api/v0/add",
                files={"file": f},
                timeout=30
            )
        response.raise_for_status()
        result = response.json()
        cid = result["Hash"]

        # 提取文本用于预览
        preview_text, file_type = extract_text_from_file(file.name)
        uploaded_files[cid] = (file.name, datetime.now().strftime("%Y-%m-%d %H:%M"), file_type)

        preview = preview_text[:1000] + ("\n...[预览截断]" if len(preview_text) > 1000 else "")
        return f"✅ 上传成功！\nCID: {cid}\n文件名: {file.name}\n类型: {file_type.upper()}\n预览:\n{preview}"
    
    except Exception as e:
        return f"❌ 上传失败: {str(e)}"

# ==================== IPFS 原始数据读取 ====================
def read_ipfs_file_raw(cid: str) -> Tuple[Optional[bytes], str]:
    try:
        response = requests.post(
            f"{IPFS_API_URL}/api/v0/cat?arg={cid}",
            timeout=30
        )
        response.raise_for_status()
        return response.content, "ok"
    except Exception as e:
        return None, f"⚠️ 文件读取失败: {str(e)}"

# ==================== 消息处理 ====================
def process_message(message: str, chat_history: list) -> Tuple[str, list, str]:
    if not message:
        return "", chat_history, "⚠️ 请输入消息"
    
    cid, status_msg = match_cid(message)
    display_message = message

    if cid:
        display_message = message.replace(cid, "上述文件", 1)
        filename, _, file_type = uploaded_files[cid]
        status_msg = f"✅ 正在分析 [{file_type.upper()}] 文件: {filename[:15]}..."

        binary_data, err = read_ipfs_file_raw(cid)
        if binary_data is None:
            return "", chat_history, err

        try:
            if file_type == "pdf":
                reader = PdfReader(BytesIO(binary_data))
                file_content = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            elif file_type == "docx":
                doc = docx.Document(BytesIO(binary_data))
                file_content = "\n".join(p.text for p in doc.paragraphs).strip()
            elif file_type == "txt":
                file_content = binary_data.decode("utf-8", errors="ignore")
            else:
                return "", chat_history, f"⚠️ 暂不支持的文件类型: {file_type}"
        except Exception as e:
            return "", chat_history, f"⚠️ 文件解析失败: {str(e)}"

        api_message = f"文件内容：\n{file_content}\n\n用户问题：{display_message}"
    else:
        api_message = message

    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json; charset=utf-8"
        }
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": api_message}],
            "temperature": 0.7
        }
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        response.raise_for_status()
        ai_response = response.json()["choices"][0]["message"]["content"]
        chat_history.append((display_message, ai_response))
    except Exception as e:
        chat_history.append((display_message, f"⚠️ API错误: {str(e)}"))

    return "", chat_history, status_msg

# ==================== Gradio UI ====================
with gr.Blocks(title="IPFS文件分析") as app:
    gr.Markdown("## 📡 IPFS文件分析系统")
    
    with gr.Row():
        file_upload = gr.UploadButton("📤 上传文件", file_types=[".txt", ".pdf", ".docx"])
        upload_status = gr.Markdown("等待文件上传...")
    
    chatbot = gr.Chatbot(height=400)
    msg = gr.Textbox(label="输入消息", placeholder="输入问题或CID...")
    status_display = gr.Markdown("🟡 等待输入...")
    
    with gr.Row():
        submit_btn = gr.Button("🚀 发送")
        clear_btn = gr.ClearButton([msg, chatbot, upload_status])

    file_upload.upload(
        handle_upload,
        inputs=file_upload,
        outputs=upload_status
    )
    
    submit_btn.click(
        process_message,
        inputs=[msg, chatbot],
        outputs=[msg, chatbot, status_display]
    )
    
    msg.change(
        fn=debounced_match_cid,
        inputs=msg,
        outputs=status_display
    )

if __name__ == "__main__":
    app.launch(server_port=7860)
